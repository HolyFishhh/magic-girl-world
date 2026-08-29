from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def paragraph_bottom_border(paragraph, color=BLUE, size="12", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_paragraph_shading(paragraph, fill):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = paragraph_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(https?://[^\s）)\]。；，]+|\*\*.+?\*\*|`.+?`|\[S\d+(?:[–-]S\d+)?\])")
SOURCE_RE = re.compile(r"^(\[S\d+\])\s+(.+?)\s*$")
URL_RE = re.compile(r"https?://[^\s；;]+")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKGREL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_TYPE_FOOTNOTES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
REL_TYPE_HYPERLINK = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
CT_FOOTNOTES = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"


def add_inline(paragraph, text, default_size=11, source_ledger=None, citations=None):
    cursor = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run, size=default_size, color=INK)
        token = match.group(0)
        if token.startswith("http"):
            add_hyperlink(paragraph, token, token)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=default_size, color=INK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=default_size - 0.5, color=DARK_BLUE)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_GRAY)
            run._element.get_or_add_rPr().append(shading)
        else:
            source_text = source_ledger.get(token) if source_ledger is not None else None
            range_match = re.fullmatch(r"\[S(\d+)[–-]S(\d+)\]", token)
            if source_text is None and source_ledger is not None and range_match:
                first, last = (int(range_match.group(1)), int(range_match.group(2)))
                source_text = "；".join(
                    source_ledger[f"[S{index}]"]
                    for index in range(first, last + 1)
                    if f"[S{index}]" in source_ledger
                )
            if source_ledger is not None and citations is not None and source_text:
                marker = f"[[FN{len(citations) + 1:03d}]]"
                run = paragraph.add_run(marker)
                set_font(run, size=default_size, color=DARK_BLUE, bold=True)
                citations.append((marker, source_text))
            else:
                run = paragraph.add_run(token)
                set_font(run, size=default_size, color=DARK_BLUE, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=default_size, color=INK)


def add_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def make(kind, abstract_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        level.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        level.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        abstract.append(level)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        return abstract, num

    bullet_abstract, bullet_num = make("bullet", next_abs, next_num)
    decimal_abstract, decimal_num = make("decimal", next_abs + 1, next_num + 1)
    numbering.append(bullet_abstract)
    numbering.append(decimal_abstract)
    numbering.append(bullet_num)
    numbering.append(decimal_num)
    return next_num, next_num + 1


def assign_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    suffix = paragraph.add_run(" 页")
    set_font(suffix, size=9, color=MUTED)


def setup_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("魔法少女世界 · 平衡与创意设计研究")
    set_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_block(document):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("剧情模式卡牌与敌人平衡创意辅助器")
    set_font(run, size=23, color="000000", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("研究、架构、实现与校准报告")
    set_font(run, size=14, color="373737")
    for label, value in (
        ("面向", "《魔法少女世界》剧情模式与第二轮变量模型"),
        ("研究日期", "2026-08-29"),
        ("研究范围", "SillyTavern、酒馆助手接口、卡牌构筑设计、AI 玩家与 PCG"),
        ("实现状态", "辅助器已接入；自动化通过；本地服务可用，内置浏览器受回环地址策略限制"),
    ):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        label_run = p.add_run(f"{label}：")
        set_font(label_run, size=11, bold=True, color="000000")
        value_run = p.add_run(value)
        set_font(value_run, size=11, color="000000")
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(14)
    paragraph_bottom_border(rule)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(p, LIGHT_BLUE)
    run = p.add_run("结论：程序负责合法性裁判与风险导演，AI 保留题材和机制创作权；机械指纹、影子模拟与真实战斗反馈共同提高挑战性、构筑关联和长期多样性。")
    set_font(run, size=11, color=DARK_BLUE, bold=True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code_block(document, lines):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.5, color=INK)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _split_markdown_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(document, rows, source_ledger, citations):
    column_count = len(rows[0])
    if column_count == 2:
        widths = [3200, 6160]
    elif column_count == 3:
        widths = [1600, 4600, 3160]
    else:
        base = CONTENT_DXA // column_count
        widths = [base] * column_count
        widths[-1] += CONTENT_DXA - sum(widths)

    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader")
    header_repeat.set(qn("w:val"), "true")
    header_properties.append(header_repeat)

    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.10
            add_inline(
                paragraph,
                value,
                default_size=9.5,
                source_ledger=source_ledger,
                citations=citations,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_markdown(document, source):
    if "## 来源" in source:
        body_source, ledger_source = source.split("## 来源", 1)
    else:
        body_source, ledger_source = source, ""
    source_ledger = {}
    for raw in ledger_source.splitlines():
        match = SOURCE_RE.match(raw.strip())
        if match:
            source_ledger[match.group(1)] = match.group(2).strip()
    citations = []
    bullet_id, decimal_id = add_numbering(document)
    lines = body_source.splitlines()
    first_section = next((i for i, value in enumerate(lines) if value.startswith("## ")), 0)
    lines = lines[first_section:]
    in_code = False
    code_lines = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if (
            line.lstrip().startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1])
        ):
            table_rows = [_split_markdown_row(line)]
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                values = _split_markdown_row(lines[index])
                if len(values) == len(table_rows[0]):
                    table_rows.append(values)
                index += 1
            add_markdown_table(document, table_rows, source_ledger, citations)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline(
                paragraph,
                heading.group(2),
                default_size={1: 16, 2: 13, 3: 12}[level],
                source_ledger=source_ledger,
                citations=citations,
            )
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            assign_numbering(paragraph, bullet_id)
            add_inline(paragraph, bullet.group(1), source_ledger=source_ledger, citations=citations)
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            assign_numbering(paragraph, decimal_id)
            add_inline(paragraph, numbered.group(1), source_ledger=source_ledger, citations=citations)
            index += 1
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.10
        add_inline(paragraph, line, source_ledger=source_ledger, citations=citations)
        index += 1
    if code_lines:
        add_code_block(document, code_lines)
    return citations


def _xml_bytes(root):
    from lxml import etree

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _next_rid(root):
    maximum = 0
    for relationship in root.findall(f"{{{PKGREL_NS}}}Relationship"):
        match = re.match(r"rId(\d+)$", relationship.get("Id") or "")
        if match:
            maximum = max(maximum, int(match.group(1)))
    return f"rId{maximum + 1}"


def _append_text_run(parent, text, style=None):
    from lxml import etree

    run = etree.SubElement(parent, f"{{{W_NS}}}r")
    if style:
        run_properties = etree.SubElement(run, f"{{{W_NS}}}rPr")
        run_style = etree.SubElement(run_properties, f"{{{W_NS}}}rStyle")
        run_style.set(f"{{{W_NS}}}val", style)
    node = etree.SubElement(run, f"{{{W_NS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    node.text = text
    return run


def inject_footnotes(docx_path, citations):
    if not citations:
        return
    from lxml import etree

    with zipfile.ZipFile(docx_path, "r") as archive:
        files = {item.filename: archive.read(item.filename) for item in archive.infolist()}

    document = etree.fromstring(files["word/document.xml"])
    document_relationships = etree.fromstring(files["word/_rels/document.xml.rels"])
    content_types = etree.fromstring(files["[Content_Types].xml"])

    footnotes = etree.Element(f"{{{W_NS}}}footnotes", nsmap={"w": W_NS, "r": R_NS})
    for note_id, separator in ((-1, "separator"), (0, "continuationSeparator")):
        note = etree.SubElement(footnotes, f"{{{W_NS}}}footnote")
        note.set(f"{{{W_NS}}}id", str(note_id))
        paragraph = etree.SubElement(note, f"{{{W_NS}}}p")
        run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
        etree.SubElement(run, f"{{{W_NS}}}{separator}")

    footnote_relationships = etree.Element(f"{{{PKGREL_NS}}}Relationships", nsmap={None: PKGREL_NS})
    namespace = {"w": W_NS}

    for note_id, (marker, note_text) in enumerate(citations, start=1):
        inserted = False
        for text_node in document.xpath(".//w:t", namespaces=namespace):
            if not text_node.text or marker not in text_node.text:
                continue
            text_node.text = text_node.text.replace(marker, "")
            run = text_node.getparent()
            while run is not None and run.tag != f"{{{W_NS}}}r":
                run = run.getparent()
            parent = run.getparent()
            reference_run = etree.Element(f"{{{W_NS}}}r")
            run_properties = etree.SubElement(reference_run, f"{{{W_NS}}}rPr")
            run_style = etree.SubElement(run_properties, f"{{{W_NS}}}rStyle")
            run_style.set(f"{{{W_NS}}}val", "FootnoteReference")
            reference = etree.SubElement(reference_run, f"{{{W_NS}}}footnoteReference")
            reference.set(f"{{{W_NS}}}id", str(note_id))
            parent.insert(parent.index(run) + 1, reference_run)
            inserted = True
            break
        if not inserted:
            raise RuntimeError(f"footnote marker missing: {marker}")

        note = etree.SubElement(footnotes, f"{{{W_NS}}}footnote")
        note.set(f"{{{W_NS}}}id", str(note_id))
        paragraph = etree.SubElement(note, f"{{{W_NS}}}p")
        paragraph_properties = etree.SubElement(paragraph, f"{{{W_NS}}}pPr")
        paragraph_style = etree.SubElement(paragraph_properties, f"{{{W_NS}}}pStyle")
        paragraph_style.set(f"{{{W_NS}}}val", "FootnoteText")
        reference_run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
        run_properties = etree.SubElement(reference_run, f"{{{W_NS}}}rPr")
        run_style = etree.SubElement(run_properties, f"{{{W_NS}}}rStyle")
        run_style.set(f"{{{W_NS}}}val", "FootnoteReference")
        etree.SubElement(reference_run, f"{{{W_NS}}}footnoteRef")
        _append_text_run(paragraph, " ")

        cursor = 0
        for match in URL_RE.finditer(note_text):
            if match.start() > cursor:
                _append_text_run(paragraph, note_text[cursor : match.start()])
            matched_url = match.group(0)
            url = matched_url.rstrip(".）)")
            trailing = matched_url[len(url) :]
            relationship_id = _next_rid(footnote_relationships)
            relationship = etree.SubElement(footnote_relationships, f"{{{PKGREL_NS}}}Relationship")
            relationship.set("Id", relationship_id)
            relationship.set("Type", REL_TYPE_HYPERLINK)
            relationship.set("Target", url)
            relationship.set("TargetMode", "External")
            hyperlink = etree.SubElement(paragraph, f"{{{W_NS}}}hyperlink")
            hyperlink.set(f"{{{R_NS}}}id", relationship_id)
            link_run = etree.SubElement(hyperlink, f"{{{W_NS}}}r")
            link_properties = etree.SubElement(link_run, f"{{{W_NS}}}rPr")
            color = etree.SubElement(link_properties, f"{{{W_NS}}}color")
            color.set(f"{{{W_NS}}}val", BLUE)
            underline = etree.SubElement(link_properties, f"{{{W_NS}}}u")
            underline.set(f"{{{W_NS}}}val", "single")
            link_text = etree.SubElement(link_run, f"{{{W_NS}}}t")
            link_text.text = url
            if trailing:
                _append_text_run(paragraph, trailing)
            cursor = match.end()
        if cursor < len(note_text):
            _append_text_run(paragraph, note_text[cursor:])

    if not any(
        relationship.get("Type") == REL_TYPE_FOOTNOTES
        for relationship in document_relationships.findall(f"{{{PKGREL_NS}}}Relationship")
    ):
        relationship = etree.SubElement(document_relationships, f"{{{PKGREL_NS}}}Relationship")
        relationship.set("Id", _next_rid(document_relationships))
        relationship.set("Type", REL_TYPE_FOOTNOTES)
        relationship.set("Target", "footnotes.xml")

    if not any(
        node.get("PartName") == "/word/footnotes.xml"
        for node in content_types.findall(f"{{{CT_NS}}}Override")
    ):
        override = etree.SubElement(content_types, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/footnotes.xml")
        override.set("ContentType", CT_FOOTNOTES)

    files["word/document.xml"] = _xml_bytes(document)
    files["word/_rels/document.xml.rels"] = _xml_bytes(document_relationships)
    files["[Content_Types].xml"] = _xml_bytes(content_types)
    files["word/footnotes.xml"] = _xml_bytes(footnotes)
    files["word/_rels/footnotes.xml.rels"] = _xml_bytes(footnote_relationships)

    temporary_path = docx_path.with_suffix(".footnotes.docx")
    with zipfile.ZipFile(temporary_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)
    shutil.move(temporary_path, docx_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: create-card-game-research-docx.py SOURCE.md OUTPUT.docx")
    source_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    source = source_path.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    setup_page(document)
    add_title_block(document)
    citations = parse_markdown(document, source)
    document.core_properties.title = "剧情模式卡牌与敌人平衡创意辅助器研究报告"
    document.core_properties.subject = "《魔法少女世界》平衡与创意辅助器的研究、实现与校准"
    document.core_properties.author = "Codex"
    document.core_properties.keywords = "卡牌构筑, 敌人设计, 平衡辅助器, 影子模拟, SillyTavern, MVU"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    inject_footnotes(output_path, citations)
    print(output_path)


if __name__ == "__main__":
    main()
