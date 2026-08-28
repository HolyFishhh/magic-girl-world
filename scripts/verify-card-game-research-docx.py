from __future__ import annotations

import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


def require(condition, message):
    if not condition:
        raise AssertionError(message)


def main():
    if len(sys.argv) != 2:
        raise SystemExit("usage: verify-card-game-research-docx.py REPORT.docx")
    path = Path(sys.argv[1]).resolve()
    document = Document(path)
    section = document.sections[0]
    require(section.page_width.twips == 12240, f"unexpected page width: {section.page_width.twips}")
    require(section.page_height.twips == 15840, f"unexpected page height: {section.page_height.twips}")
    for value, label in (
        (section.top_margin, "top margin"),
        (section.right_margin, "right margin"),
        (section.bottom_margin, "bottom margin"),
        (section.left_margin, "left margin"),
    ):
        require(value.twips == 1440, f"unexpected {label}: {value.twips}")
    require(abs(section.header_distance.twips - 708) <= 1, f"unexpected header distance: {section.header_distance.twips}")
    require(abs(section.footer_distance.twips - 708) <= 1, f"unexpected footer distance: {section.footer_distance.twips}")

    expected_styles = {
        "Normal": (11, None),
        "Heading 1": (16, "2E74B5"),
        "Heading 2": (13, "2E74B5"),
        "Heading 3": (12, "1F4D78"),
    }
    for name, (size, color) in expected_styles.items():
        style = document.styles[name]
        require(style.font.size.pt == size, f"{name} size drifted")
        require(style.font.name == "Calibri", f"{name} font drifted")
        if color:
            require(str(style.font.color.rgb) == color, f"{name} color drifted")

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for required in (
        "卡牌与敌人设计知识库研究报告",
        "《杀戮尖塔2》仍为 Early Access",
        "来源台账",
        "在《魔法少女世界》0.5.146 的落地",
        "卡牌范式 499 token",
        "敌人范式 543 token",
    ):
        require(required in text, f"missing report content: {required}")
    for source_id in range(1, 25):
        require(f"[S{source_id}]" in text, f"missing source marker S{source_id}")

    with ZipFile(path, "r") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    require(document_xml.count("<w:hyperlink") >= 24, "source URLs must be clickable hyperlinks")
    require(relationships.count('TargetMode="External"') >= 24, "source hyperlink relationships are incomplete")
    require('w:left="540"' in numbering_xml, "list text indent must be 0.375in")
    require('w:hanging="270"' in numbering_xml, "list hanging indent must be about 0.188in")
    require('w:after="80"' in numbering_xml, "list spacing after must be 4pt")
    require('w:line="300"' in numbering_xml, "list line spacing must be 1.25")

    for table_index, table in enumerate(document.tables):
        tbl_pr = table._tbl.tblPr
        width = tbl_pr.find(qn("w:tblW"))
        indent = tbl_pr.find(qn("w:tblInd"))
        require(width is not None and width.get(qn("w:w")) == "9360", f"table {table_index} width drifted")
        require(indent is not None and indent.get(qn("w:w")) == "120", f"table {table_index} indent drifted")
        grid_total = sum(int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.findall(qn("w:gridCol")))
        require(grid_total == 9360, f"table {table_index} grid width drifted: {grid_total}")

    heading_counts = {
        name: sum(1 for paragraph in document.paragraphs if paragraph.style.name == name)
        for name in ("Heading 1", "Heading 2", "Heading 3")
    }
    require(heading_counts["Heading 1"] >= 10, "report section hierarchy is incomplete")
    require(len(document.paragraphs) >= 250, "report appears truncated")
    require(len(document.tables) >= 4, "callout and code-block tables are incomplete")
    print(
        f"DOCX audit passed: paragraphs={len(document.paragraphs)}, tables={len(document.tables)}, "
        f"headings={heading_counts}, hyperlinks={document_xml.count('<w:hyperlink')}."
    )


if __name__ == "__main__":
    main()
