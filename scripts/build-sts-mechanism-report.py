from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "research" / "杀戮尖塔机制覆盖评估-source.md"
OUTPUT = ROOT / "docs" / "research" / "杀戮尖塔机制覆盖评估.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
BORDER = "B8C5D3"
CALLOUT_FILL = "F4F6F9"
INK = "17212B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def set_fixed_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=INK, bold=False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id + 1000:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_inline(paragraph, text: str, *, size=None, color=INK, bold=False) -> None:
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=DARK_BLUE, bold=False)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color, bold=bold)


def configure_document(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Table Citation" not in doc.styles:
        style = doc.styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Table Citation"]
    set_style_font(style, 9, MUTED)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.0

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("机制覆盖评估")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = p.add_run("\t2026-08-29")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = fp.add_run("第 ")
    set_run_font(prefix, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)
    suffix = fp.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)

    return add_numbering(doc, "bullet"), add_numbering(doc, "decimal")


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("ENGINEERING RESEARCH BRIEF")
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    add_inline(title, "《杀戮尖塔 1/2》机制覆盖评估", size=23, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    add_inline(subtitle, "通用战斗 DSL、当前实现边界与后续路线图", size=13, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    values = [
        ("更新时间", "2026-08-29"),
        ("研究范围", "本项目、STS1 正式版、STS2 当前公开 EA/Beta"),
        ("版本边界", "工程规划稿；不代表 STS2 未来正式版最终机制"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], TABLE_FILL)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for idx, cell in enumerate(row.cells):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=9.5, color=INK, bold=(idx == 0))
    set_fixed_table_geometry(meta, [1700, 7660])
    set_table_borders(meta)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    widths = [1650, 2650, 3700, 1360] if column_count == 4 else [9360 // column_count] * column_count
    widths[-1] += 9360 - sum(widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, value, size=8.5, color=INK, bold=(row_index == 0))
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
    set_fixed_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph(style="Table Citation")
    after.add_run("状态依据：截至 2026-08-29 的代码审计与公开资料；“部分具备”不等于完整语义覆盖。")


def add_list_paragraph(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    add_inline(p, text)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    bullet_id, decimal_id = configure_document(doc)
    add_title_block(doc)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if i < 5 or not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[4:], size=12, color=DARK_BLUE, bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[3:], size=13, color=BLUE, bold=True)
        elif line.startswith("# "):
            pass
        elif re.match(r"^\d+\.\s+", line):
            add_list_paragraph(doc, re.sub(r"^\d+\.\s+", "", line), decimal_id)
        elif line.startswith("- "):
            add_list_paragraph(doc, line[2:], bullet_id)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    audit_output()
    print(OUTPUT)


def audit_output() -> None:
    """Fail the build when the DOCX drifts from the selected compact-reference preset."""
    doc = Document(OUTPUT)
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert section.left_margin == Inches(1)
    assert abs(section.header_distance - Inches(0.492)) <= 1000
    assert abs(section.footer_distance - Inches(0.492)) <= 1000

    expected_styles = {
        "Normal": (11, 0, 6, 1.25),
        "Heading 1": (16, 18, 10, 1.0),
        "Heading 2": (13, 14, 7, 1.0),
        "Heading 3": (12, 10, 5, 1.0),
    }
    for name, (size, before, after, spacing) in expected_styles.items():
        style = doc.styles[name]
        assert style.font.size == Pt(size), name
        assert style.paragraph_format.space_before == Pt(before), name
        assert style.paragraph_format.space_after == Pt(after), name
        assert style.paragraph_format.line_spacing == spacing, name

    assert len(doc.tables) == 2
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "9360"
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == "120"
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360
        for row in table.rows:
            cell_widths = [int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w"))) for cell in row.cells]
            assert cell_widths == grid_widths

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "执行摘要" in text
    assert "机制覆盖矩阵" in text
    assert "所有网页访问日期：2026-08-29" in text
    assert "<placeholder>" not in text.lower()
    print("DOCX preset and structure audit passed.")


if __name__ == "__main__":
    build()
